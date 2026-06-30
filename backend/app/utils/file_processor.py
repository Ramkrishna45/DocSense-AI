"""File and URL processing utilities for extracting text from various sources."""

import logging
import re
from pathlib import Path
import requests
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> list[tuple[int, str]]:
    """Extract text from a PDF file, page by page."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    pages = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            if text.strip():
                pages.append((page_num + 1, text.strip()))
        doc.close()
    except Exception as e:
        logger.error("Failed to extract text from PDF %s: %s", file_path, str(e))
        raise ValueError(f"Failed to process PDF file: {str(e)}") from e

    return pages

def extract_text_from_docx(file_path: str) -> list[tuple[int | None, str]]:
    """Extract text from a DOCX file."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = DocxDocument(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        if full_text.strip():
            return [(None, full_text)]
        return []
    except Exception as e:
        logger.error("Failed to extract text from DOCX %s: %s", file_path, str(e))
        raise ValueError(f"Failed to process DOCX file: {str(e)}") from e

def extract_text_from_txt(file_path: str) -> list[tuple[int | None, str]]:
    """Extract text from a plain text or markdown file."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Text file not found: {file_path}")

    try:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")

        if text.strip():
            return [(None, text.strip())]
        return []
    except Exception as e:
        logger.error("Failed to read text file %s: %s", file_path, str(e))
        raise ValueError(f"Failed to process text file: {str(e)}") from e

def extract_text_from_youtube(url: str) -> list[tuple[int | None, str]]:
    """Extract transcript from YouTube video URL."""
    try:
        video_id = None
        if "youtu.be/" in url:
            video_id = url.split("youtu.be/")[1].split("?")[0]
        elif "v=" in url:
            video_id = url.split("v=")[1].split("&")[0]
            
        if not video_id:
            raise ValueError(f"Could not extract video ID from {url}")
            
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        
        # We can store timestamp blocks, but for now we concatenate all text
        full_text = " ".join([entry['text'] for entry in transcript])
            
        return [(None, full_text)]
    except Exception as e:
        logger.error(f"Failed to extract YouTube transcript from {url}: {e}")
        raise ValueError(f"Failed to extract YouTube transcript: {e}")

def extract_text_from_url(url: str) -> list[tuple[int | None, str]]:
    """Extract main text content from a general URL or GitHub raw URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        
        if "github.com" in url and "/blob/" in url:
            url = url.replace("github.com", "raw.githubusercontent.com").replace("/blob/", "/")
            
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        content_type = response.headers.get('Content-Type', '').lower()
        if 'text/plain' in content_type or url.endswith('.md'):
            return [(None, response.text)]
            
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove non-content elements
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.extract()
            
        text = soup.get_text(separator=' ', strip=True)
        return [(None, text)]
    except Exception as e:
        logger.error(f"Failed to extract text from URL {url}: {e}")
        raise ValueError(f"Failed to extract URL content: {e}")

def process_source(source: str, source_type: str = 'file', original_filename: str = '') -> list[tuple[int | None, str]]:
    """Process a source and extract text based on its type.
    
    Args:
        source: Path to file or URL string.
        source_type: 'file', 'url', 'youtube', 'github'
        original_filename: needed for file extraction to know extension
    """
    if source_type == 'youtube':
        return extract_text_from_youtube(source)
    elif source_type in ['url', 'github']:
        return extract_text_from_url(source)
    else:
        file_type = original_filename.split('.')[-1].lower() if original_filename else source.split('.')[-1].lower()
        extractors = {
            "pdf": extract_text_from_pdf,
            "docx": extract_text_from_docx,
            "txt": extract_text_from_txt,
            "md": extract_text_from_txt,
        }
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        return extractor(source)
